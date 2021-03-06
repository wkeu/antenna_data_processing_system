from docx import Document
from docx.shared import Inches
import pandas as pd
import os

def panda_to_word_table(doc,df,measurment_type="El tilt deviation",spec="",style="Light List Accent 1"):
    
    # Function to write excel spreadsheet to word table
    #Formatting of panda dataframe
    df = df.rename(index=str, columns={"index": "Frequency"})
        
    ###############################################################################
    # Data table (1), just copy all the data of the panda. Port names, frequency etc.
    ###############################################################################
    # add a table to the end and create a reference antenna_type
    # extra row is so we can add the header row
    t = doc.add_table(df.shape[0]+1, df.shape[1]+2)
    t.style =style
    
    # add the header rows.
    for j in range(2,df.shape[-1]+2):
        t.cell(0,j).text = df.columns[j-2]
    
    # add the rest of the data frame
    for i in range(1,df.shape[0]+1):
        for j in range(df.shape[-1]):
            t.cell(i,j+2).text = str(df.values[i-1,j])
    
    ###############################################################################
    # Info Table I (2), Item and spec
    ###############################################################################
    t.cell(0,0).text="Item"
    t.cell(1,0).text=measurment_type
    
    #Merge cells
    a = t.cell(1, 0)
    b = t.cell(df.shape[0],0)#
    a.merge(b)
    
    t.cell(0,1).text="Spec"
    t.cell(1,1).text=spec
    
    #Merge cells
    a = t.cell(1, 1)
    b = t.cell(df.shape[0],1)#
    a.merge(b)
    
    ###############################################################################
    # Info Table II (3), mean, max, min, pass
    ###############################################################################
    
    df = df.set_index('Frequency')
    
    t = doc.add_table(4, 2)
    t.style = style
    
    cell = t.cell(0, 0)
    cell.text = 'Average'
    mean=df.values.mean()
    t.cell(0, 1).text=str(round(mean,2))
    
    cell = t.cell(1, 0)
    cell.text = 'Max'
    t.cell(1, 1).text=str(df.values.max())
    
    cell = t.cell(2, 0)
    cell.text = 'Min'
    t.cell(2, 1).text=str(df.values.min())
    
    cell = t.cell(3, 0)
    cell.text = 'Result'
    midpoint = ((df.values.max() - df.values.min())/2)+df.values.min()
    tolerance = ((df.values.max() - df.values.min())/2)
    t.cell(3, 1).text=str(round(midpoint)) +' ± ' + str(round(tolerance))
    doc.add_page_break()
    
#Function to place images into report
def insert_cart_plots(doc,images_dir_path):

    file_list=list()
    
    doc.add_heading("Plots", level=1)
    
    #Create a list of all image files
    for path, subdirs, files in os.walk(images_dir_path):
           for filename in files:
             f = os.path.join(path, filename)
             file_list.append(str(f)) 
    
    #Put Image into document
    for picture in file_list:
        #heading
        name=picture.split("\\")[-1]
        name=name.split(".")[0]
        
        doc.add_heading(name, level=5)
        doc.add_picture(picture,width=Inches(6.6))

def insert_radiation_tables(doc,master_table_path):

    #Open excel document
    xl = pd.ExcelFile(master_table_path)
    #Load in all sheet names
    sheet_names = xl.sheet_names
    
    for sheet_name in sheet_names:
        doc.add_heading(sheet_name, level=5)
        df = xl.parse(sheet_name)
        panda_to_word_table(doc,df,measurment_type=sheet_name,spec="")   
        
     
        
        
        
        
        

###############################################################################
#
# Main
#
###############################################################################        

def generate_report(dir_path,antenna_model="AWXXXX"):
    
    master_table_path = dir_path + "master_table.xlsx"
    gain_table_path = dir_path + "gain_table.xlsx"
    images_dir_path=dir_path + "images/CART"
    gain_images_dir_path=dir_path + "images/GAIN"
    save_path=dir_path +antenna_model+"_report.docx"
    
    # Open a new document
    doc = Document()
    doc.add_heading('Antenna Report', 0)
    
    
    #doc.add_heading("Electrical Tables", level=1)
   #insert_gain_tables(doc,gain_table_path)
    
    #doc.add_page_break()
    doc.add_heading("Radiation Tables", level=1)
    
    insert_radiation_tables(doc,master_table_path)
   
    #Add the plots
    insert_cart_plots(doc,images_dir_path)
    
    
    # save the doc
    doc.save(save_path)
    print("Report for antenna " + antenna_model + " generated.")